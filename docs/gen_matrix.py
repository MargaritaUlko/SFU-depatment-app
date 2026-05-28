"""Generate permissions matrix as .docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROLES = ["student", "headman", "teacher", "dean", "deputy_head", "admin"]
ROLE_RU = {
    "student": "Студент",
    "headman": "Староста",
    "teacher": "Преподаватель",
    "dean": "Деканат",
    "deputy_head": "Зав. каф.",
    "admin": "Администратор",
}

YES = "✅"
NO = "❌"
OWN = "🔒"

SECTIONS = [
    {
        "title": "Пользователи (/users)",
        "rows": [
            ("Список / поиск пользователей",         YES, YES, YES, YES, YES, YES),
            ("Список преподавателей",                 YES, YES, YES, YES, YES, YES),
            ("Просмотр / загрузка аватара",           YES, YES, YES, YES, YES, YES),
            ("Обновить профиль",                      OWN, OWN, OWN, OWN, OWN, YES),
            ("Сменить пароль",                        OWN, OWN, OWN, OWN, OWN, YES),
            ("Профиль студента (/me/student-profile)",YES, YES,  NO,  NO,  NO,  NO),
            ("Профиль преподавателя (/me/teacher-profile)", NO, NO, YES, NO, YES, NO),
            ("Профиль деканата (/me/dean-profile)",    NO,  NO,  NO, YES,  NO,  NO),
            ("Назначить группу студенту",              NO,  NO,  NO, YES,  NO,  NO),
            ("Создать пользователя",                   NO,  NO,  NO,  NO,  NO, YES),
            ("Удалить пользователя",                   NO,  NO,  NO,  NO,  NO, YES),
        ],
    },
    {
        "title": "Документы (/documents)",
        "note": "🔒 Список/просмотр/скачивание фильтруется по visibility (роль / group:id / stream:id). Редактирование и удаление — только свои документы, кроме admin.",
        "rows": [
            ("Список / просмотр / скачать",            OWN, OWN, OWN, OWN, OWN, OWN),
            ("Загрузить документ",                      NO,  NO, YES, YES, YES, YES),
            ("Редактировать метаданные",                 NO,  NO, OWN, OWN, OWN, YES),
            ("Удалить",                                  NO,  NO, OWN, OWN, OWN, YES),
        ],
    },
    {
        "title": "События (/events)",
        "note": "🔒 teacher и headman — только своё событие; deputy_head, dean, admin — любое.",
        "rows": [
            ("Список / просмотр",                      YES, YES, YES, YES, YES, YES),
            ("Создать событие",                          NO, YES, YES, YES, YES, YES),
            ("Редактировать событие",                    NO, OWN, OWN, YES, YES, YES),
            ("Удалить событие",                          NO, OWN, OWN, YES, YES, YES),
            ("Загрузить картинку события",               NO, OWN, OWN, YES, YES, YES),
        ],
    },
    {
        "title": "Объявления (/announcements)",
        "note": "¹ Удаление — роутер пропускает любого авторизованного, ограничения на уровне сервиса.",
        "rows": [
            ("Список / просмотр / мои",                YES, YES, YES, YES, YES, YES),
            ("Создать",                                  NO, YES, YES,  NO, YES, YES),
            ("Редактировать (PATCH)",                    NO, YES, YES, YES, YES, YES),
            ("Архивировать",                             NO, YES,  NO, YES, YES, YES),
            ("Восстановить из архива",                   NO,  NO,  NO, YES,  NO, YES),
            ("Удалить ¹",                               OWN, OWN, OWN, OWN, OWN, OWN),
        ],
    },
    {
        "title": "Посещаемость (/attendance)",
        "rows": [
            ("Сгенерировать QR-токен",                   NO, YES, YES,  NO, YES,  NO),
            ("Получить QR-изображение",                  NO, YES, YES,  NO, YES,  NO),
            ("Отметиться по QR",                        YES, YES,  NO,  NO,  NO,  NO),
            ("Ручная отметка студента",                  NO, YES, YES,  NO, YES,  NO),
            ("Посещаемость занятия / студента",         YES, YES, YES, YES, YES, YES),
        ],
    },
    {
        "title": "Сообщения — broadcast (/messages)",
        "note": "Методы PUT/DELETE объявлены, но не реализованы.",
        "rows": [
            ("Список / отправить",                       NO, YES, YES,  NO, YES, YES),
            ("Просмотр по ID",                           NO, OWN, OWN,  NO, OWN, YES),
        ],
    },
    {
        "title": "Чаты (/chats)",
        "note": "🔒 Только участник чата; admin видит и пишет в любой чат.",
        "rows": [
            ("Список чатов",                            OWN, OWN, OWN, OWN, OWN, YES),
            ("Открыть личный чат",                      YES, YES, YES, YES, YES, YES),
            ("Создать групповой чат",                   YES, YES, YES, YES, YES, YES),
            ("Читать сообщения чата",                   OWN, OWN, OWN, OWN, OWN, YES),
            ("Отправить сообщение (HTTP + WS)",         OWN, OWN, OWN, OWN, OWN, YES),
        ],
    },
    {
        "title": "ВКР (/vkr)",
        "rows": [
            ("Предложить тему",                         YES, YES, YES,  NO, YES,  NO),
            ("Мои темы",                                YES, YES, YES, YES, YES, YES),
            ("Просмотр темы по ID",                     YES, YES, YES, YES, YES, YES),
            ("Одобренные темы",                          NO,  NO,  NO, YES, YES, YES),
            ("Все темы (любой статус)",                  NO,  NO,  NO,  NO, YES, YES),
            ("Рецензировать тему",                       NO,  NO,  NO,  NO, YES,  NO),
        ],
    },
    {
        "title": "Потоки (/streams) и Группы (/groups)",
        "rows": [
            ("Список / просмотр",                       YES, YES, YES, YES, YES, YES),
            ("Создать / редактировать / удалить",         NO,  NO,  NO, YES,  NO, YES),
        ],
    },
    {
        "title": "Расписание (/lessons)",
        "rows": [
            ("Расписание группы / преподавателя",       YES, YES, YES, YES, YES, YES),
            ("Синхронизация расписания",                  NO,  NO,  NO, YES,  NO, YES),
        ],
    },
]

COLOR_HEADER = RGBColor(0x1F, 0x49, 0x7D)   # dark blue
COLOR_SECTION = RGBColor(0xD6, 0xE4, 0xF0)  # light blue
COLOR_YES = RGBColor(0xE2, 0xEF, 0xDA)       # light green
COLOR_NO = RGBColor(0xFC, 0xE4, 0xD6)        # light red/orange
COLOR_OWN = RGBColor(0xFF, 0xF2, 0xCC)       # light yellow
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def cell_text(cell, text, bold=False, center=False, font_size=9, color=None):
    para = cell.paragraphs[0]
    para.clear()
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_section(doc: Document, section: dict):
    # Section heading
    p = doc.add_paragraph()
    p.style = "Heading 2"
    run = p.add_run(section["title"])
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_HEADER

    # Optional note
    if "note" in section:
        np = doc.add_paragraph(section["note"])
        np.runs[0].font.size = Pt(8)
        np.runs[0].font.italic = True

    cols = 1 + len(ROLES)
    rows_data = section["rows"]
    table = doc.add_table(rows=1 + len(rows_data), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Column widths
    widths = [Cm(7.5)] + [Cm(2.2)] * len(ROLES)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

    # Header row
    header = table.rows[0]
    set_cell_bg(header.cells[0], COLOR_HEADER)
    cell_text(header.cells[0], "Действие", bold=True, font_size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
    for i, role in enumerate(ROLES):
        set_cell_bg(header.cells[i + 1], COLOR_HEADER)
        cell_text(header.cells[i + 1], ROLE_RU[role], bold=True, center=True, font_size=8,
                  color=RGBColor(0xFF, 0xFF, 0xFF))

    # Data rows
    for r_idx, row_data in enumerate(rows_data):
        row = table.rows[r_idx + 1]
        action = row_data[0]
        values = row_data[1:]
        bg = COLOR_SECTION if r_idx % 2 == 0 else COLOR_WHITE
        set_cell_bg(row.cells[0], bg)
        cell_text(row.cells[0], action, font_size=9)
        for c_idx, val in enumerate(values):
            if val == YES:
                set_cell_bg(row.cells[c_idx + 1], COLOR_YES)
            elif val == NO:
                set_cell_bg(row.cells[c_idx + 1], COLOR_NO)
            elif val == OWN:
                set_cell_bg(row.cells[c_idx + 1], COLOR_OWN)
            else:
                set_cell_bg(row.cells[c_idx + 1], bg)
            cell_text(row.cells[c_idx + 1], val, center=True, font_size=11)

    doc.add_paragraph()  # spacing


def build():
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Матрица прав доступа по ролям")
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = COLOR_HEADER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = subtitle.add_run("Информационный портал кафедры — REST API")
    rs.font.size = Pt(10)
    rs.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    doc.add_paragraph()

    # Legend
    leg = doc.add_paragraph()
    leg.add_run("Обозначения:  ").font.bold = True
    leg.add_run("✅ — полный доступ   🔒 — ограниченный (только своё / только участник)   ❌ — нет доступа")
    leg.runs[-1].font.size = Pt(9)
    doc.add_paragraph()

    for s in SECTIONS:
        add_section(doc, s)

    out = r"C:\Users\kogne\Desktop\department-app\docs\permissions_matrix.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build()
